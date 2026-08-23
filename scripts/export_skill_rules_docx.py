# -*- coding: utf-8 -*-
"""Export current skill rules markdown into a formatted Word document on Desktop."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_PATH = Path.home() / "Desktop" / "超限德州-技能系统现行规则.docx"

NAVY = RGBColor(0x1A, 0x24, 0x3A)
GOLD = RGBColor(0x8A, 0x6A, 0x2F)
RED = RGBColor(0x8B, 0x1E, 0x2D)
MUTED = RGBColor(0x4A, 0x52, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1A243A"
TABLE_ALT_BG = "F4F1EA"
TABLE_BORDER = "C8BFA8"


def set_run_font(run, name="微软雅黑", size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=6, line=18, exact=False):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    if exact:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def shade_cell(cell, hex_color):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color=TABLE_BORDER):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def write_cell(cell, text, *, header=False, center=False, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center or header else WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=0, line=16, exact=True)
    run = p.add_run(str(text))
    set_run_font(
        run,
        size=size,
        bold=header,
        color=WHITE if header else NAVY,
    )
    shade_cell(cell, TABLE_HEADER_BG if header else "FFFFFF")
    set_cell_borders(cell)
    set_cell_margins(cell)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    cell._tc.get_or_add_tcPr().append(vAlign)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)

    for i, header in enumerate(headers):
        write_cell(table.rows[0].cells[i], header, header=True, center=True, size=9)
    prevent_row_split(table.rows[0])

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            center = c_idx == 0 or (len(headers) <= 4 and c_idx < len(headers) - 1)
            write_cell(table.rows[r_idx + 1].cells[c_idx], value, center=center, size=9)
            if r_idx % 2 == 1:
                shade_cell(table.rows[r_idx + 1].cells[c_idx], TABLE_ALT_BG)
        prevent_row_split(table.rows[r_idx + 1])

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=8, line=6, exact=True)
    return table


def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph_spacing(p, before=16, after=8, line=22)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, size=16, bold=True, color=NAVY)
        # gold underline via bottom border
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "8A6A2F")
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
    elif level == 2:
        set_paragraph_spacing(p, before=12, after=6, line=20)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True, color=RED)
    else:
        set_paragraph_spacing(p, before=8, after=4, line=18)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_run_font(run, size=11.5, bold=True, color=GOLD)
    return p


def add_body(doc, text, *, bold=False, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=after, line=18)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold, color=NAVY)
    return p


def add_bullet(doc, text, *, level=0, number=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=1, after=2, line=17)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    prefix = f"{number}. " if number is not None else "• "
    run = p.add_run(prefix + text)
    set_run_font(run, size=10.5, color=NAVY)
    return p


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ""
    shade_cell(cell, "F7F1E3")
    set_cell_borders(cell, "8A6A2F")
    set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=0, line=18)
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=True, color=GOLD)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=8, line=6, exact=True)


def add_skill_meta(doc, line):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=16)
    run = p.add_run(line)
    set_run_font(run, size=10, italic=True, color=MUTED)


def set_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(hp, before=0, after=0, line=14)
    run = hp.add_run("超限德州  /  OVERLIMIT: HOLD'EM    技能系统现行规则")
    set_run_font(run, size=8, color=GOLD)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(fp, before=0, after=0, line=14)
    run = fp.add_run("实现版基线  ·  与当前代码一致  ·  2026-08-22  ·  第 ")
    set_run_font(run, size=8, color=MUTED)
    # PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r1 = fp.add_run()
    r1._r.append(fldChar1)
    r2 = fp.add_run()
    r2._r.append(instr)
    r3 = fp.add_run()
    r3._r.append(fldChar2)
    for r in (r1, r2, r3):
        set_run_font(r, size=8, color=MUTED)
    run = fp.add_run(" 页")
    set_run_font(run, size=8, color=MUTED)


def add_cover(doc):
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=18)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(kicker, before=0, after=6, line=18)
    run = kicker.add_run("OVERLIMIT: HOLD'EM")
    set_run_font(run, size=12, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=4, after=8, line=32)
    run = title.add_run("超限德州")
    set_run_font(run, size=28, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(sub, before=0, after=18, line=22)
    run = sub.add_run("技能系统现行规则（实现版）")
    set_run_font(run, size=16, bold=True, color=RED)

    add_callout(doc, "本文档是与当前代码一致的规则基线，不是重新设计稿。感知与强运保持 FROZEN_V1。贷款信用受限 V2 已升为正式规则，100 / 150 与公平负载费用不再因此削弱。")

    meta_rows = [
        ["文档性质", "现行实现对照，可直接作为开发 / 审阅基线"],
        ["首发规模", "24 个主体技能 + 9 个协议"],
        ["构筑约束", "最少 1 个、最多 4 个，总负载上限 8"],
        ["冻结项", "感知 FROZEN_V1 / spec-25-50；强运 FROZEN_V1 / soft-v1；贷款信用受限 V2 正式化"],
        ["冻结日期", "感知 / 强运 2026-08-20；贷款信用 2026-08-21"],
        ["对照代码", "definitions.js / skillConfig.js / skillEngine.js / skillState.js / gameEngine.js / chipEconomy.js / handRankBonus.js"],
    ]
    add_table(doc, ["项", "说明"], meta_rows, col_widths=[3.8, 12.6])

    toc_intro = doc.add_paragraph()
    toc_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(toc_intro, before=8, after=6, line=18)
    run = toc_intro.add_run("目录")
    set_run_font(run, size=13, bold=True, color=NAVY)

    toc = [
        "0. 阅读约定",
        "1. 全局规则",
        "2. 首发一览",
        "3. 主体技能（1–24）",
        "4. 协议（9）",
        "5. 关键交互摘要",
        "6. 明确不回滚的事实",
        "7. 冻结声明",
    ]
    for item in toc:
        add_bullet(doc, item)

    doc.add_page_break()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.font.color.rgb = NAVY
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    set_page(doc)
    add_cover(doc)

    # 0
    add_heading_styled(doc, "0. 阅读约定", 1)
    terms = [
        ("标准下注贡献", "盲注、跟注、加注、标准 ALL IN 投入。进入底池，可被撤退退还，可被血战 / 绝境 / 绝路 / 协议 / 防守修正。"),
        ("直接技能筹码转移", "贷款拿走的筹码、终局没收未匹配贡献等。不进倍率，撤退不退还，公平不回滚已完成转移。"),
        ("技能事件", "一次合法、已扣费（若需要）并进入结算的主动发动，或一次真正触发的被动效果。"),
        ("既定事实", "已看见的信息、已完成的换牌 / 牌堆修改、已发生的筹码转移。公平与撤退都不能时间倒流。"),
        ("公开 / 秘密", "公开技能双方都知道发生了什么；秘密技能默认只对发动者可见，满足特定揭露条件时才可被确认。"),
    ]
    add_table(
        doc,
        ["术语", "含义"],
        [[k, v] for k, v in terms],
        col_widths=[4.2, 12.2],
    )

    # 1
    add_heading_styled(doc, "1. 全局规则", 1)
    add_heading_styled(doc, "1.0 全模式基础经济与牌型奖励", 2)
    add_body(
        doc,
        "牌型基础奖励属于《超限德州》的游戏基础经济规则，不是技能效果。无论 standard + skill off、overdrive + skill off、standard + skill on、overdrive + skill on，只要最终通过 Showdown 决出胜者，就发放 launch-v1 牌型基础奖励。关闭技能只代表不启用技能系统，并不代表完全原版 Texas Hold'em。",
    )
    add_table(
        doc,
        ["牌型", "category", "奖励"],
        [
            ["高牌", "1", "+0"],
            ["一对", "2", "+0"],
            ["两对", "3", "+0"],
            ["三条", "4", "+25"],
            ["顺子", "5", "+50"],
            ["同花", "6", "+75"],
            ["葫芦", "7", "+100"],
            ["四条", "8", "+250"],
            ["同花顺", "9", "+400"],
            ["皇家同花顺", "10", "+500"],
        ],
        col_widths=[4.0, 2.4, 10.0],
    )
    for t in [
        "Fold / Retreat / 平局：不发牌型奖励。",
        "皇家同花顺是正式独立牌型（category 10）。协议 P09 同时覆盖 category 9 和 10，不另设皇家同花顺独立协议。",
        "牌型奖励不是「自己的技能倍率」，因此不阻止 Protocol。Blood / Desperation / Dead End 等已有自己的技能倍率时，仍按现有规则阻止自己的 Protocol。",
    ]:
        add_bullet(doc, t)
    add_callout(
        doc,
        "正式结算顺序：standardPokerNet + handRankBonus + 其他基础加值（例如 Probe）→ 自身合法技能倍率 → 对手产生的合法倍率 → Defense → Stack Cap → 唯一整数筹码转移",
    )

    add_heading_styled(doc, "1.1 构筑", 2)
    add_table(
        doc,
        ["项", "现行值"],
        [
            ["最少技能", "1"],
            ["最多技能", "4"],
            ["总负载上限", "8"],
            ["更换时机", "仅开局前；对局中不可更换"],
            ["对手可见构筑", "开局隐藏数量、负载与具体技能"],
        ],
        col_widths=[4.2, 12.2],
    )
    add_body(doc, "合法示例：", bold=True, after=2)
    for t in ["终局 6", "终局 6 + 警觉 1 + 深呼吸 1", "公平 4 + 反制 4", "公平 4 + 伪装 4"]:
        add_bullet(doc, t)
    add_body(doc, "非法示例：", bold=True, after=2)
    for t in [
        "空构筑",
        "公平 4 + 天命 5（负载 9）",
        "公平 4 + 千术 5",
        "公平 4 + 强运 5",
        "公平 4 + 零化 5",
        "终局 6 + 千术 5",
    ]:
        add_bullet(doc, t)

    add_heading_styled(doc, "1.2 能量", 2)
    add_table(
        doc,
        ["项", "现行值"],
        [
            ["开局能量", "4"],
            ["通常上限", "8"],
            ["携带天命时本人上限", "10"],
            ["对手界面显示上限", "8（不直接显示 10）"],
            ["强运允许下限", "-4"],
            ["负能量限制", "除强运外，不能发动主动技能，也不能触发新的被动技能事件"],
            ["普通公开快照", "clamp(realFinalEnergy, 0, 8)"],
        ],
        col_widths=[4.8, 11.6],
    )
    add_body(doc, "手牌结束自然恢复：", bold=True, after=4)
    add_table(
        doc,
        ["结果", "能量"],
        [
            ["胜", "+0"],
            ["败（含普通 Fold）", "+1"],
            ["平", "+0"],
            ["撤退 Fold", "无败者 +1"],
        ],
        col_widths=[6.0, 10.4],
    )
    add_body(doc, "公平成功后，本手全部结束恢复都被抑制，包括败者 +1、深呼吸、绝境 +1、反制空放返还、资源型强运。")
    add_body(doc, "对手能量规则：逐手公开、手内冻结。", bold=True, after=4)
    for t in [
        "每手结束：所有能量恢复、Loan 还款、债务、Fairness 抑制、Fortune 资源效果等全部处理结束后，才刷新公开快照。",
        "下一手进行过程中：对手显示保持冻结，不实时变化。",
        "普通对手可见值：真实 -4/-3/-1/0 → 0；真实 1~8 → 真实值；真实 9/10 → 8。",
        "本人始终看到自己的真实当前能量，包括 Strong Fortune 负数和 Destiny 的 9/10。",
        "灵视始终读取服务器真实当前能量（含负数、9/10 与手内实时变化），只写入私有结果，不得顶替对手信息条上的公开冻结值。",
        "普通客户端不得通过任何其他字段区分真实 0 与被遮蔽的负数 0，也不得区分真实 8 与 Destiny 被封顶后的 8。",
    ]:
        add_bullet(doc, t)

    add_heading_styled(doc, "1.3 倍率与直接筹码", 2)
    add_body(doc, "只修正标准净筹码转移。终局没收、贷款直接拿筹不进倍率。")
    add_body(doc, "常见乘法：", bold=True, after=2)
    for t in ["血战 ×2；双方血战 ×4", "绝境获胜 ×3", "绝路因对手普通 Fold ×3", "协议获胜 ×2"]:
        add_bullet(doc, t)
    add_body(doc, "叠加规则：", bold=True, after=2)
    for t in [
        "自己的血战 / 绝境 / 绝路 与 对手的血战 可乘法叠加",
        "若本手已有自己技能产生的其他筹码倍率，自己的协议不触发",
        "对手技能产生的倍率不阻止自己的协议",
        "试探 +50 先并入基础标准收益，再参与倍率",
        "最终支付不超过败方实际可支付筹码",
    ]:
        add_bullet(doc, t)

    add_heading_styled(doc, "1.4 反制、公平、绝密（总则）", 2)
    for t in [
        "主动技能默认可被反制；公平不能被反制。",
        "反制捕获的是「下一次合法且已扣费的主动技能」，不是非法请求，也不是被动技能本身。",
        "公平清除持续 / 预埋 / 待结算状态，封锁本手后续主动与新被动事件，并抑制本手结束恢复。已完成事实不回滚。",
        "绝密保护的是本人底牌私人信息，不阻挡公共牌情报、灵视、纯技能元信息。",
    ]:
        add_bullet(doc, t)

    # 2
    add_heading_styled(doc, "2. 首发一览", 1)
    add_table(
        doc,
        ["#", "技能", "负载", "费用", "类型", "可见", "每手"],
        [
            ["1", "深呼吸", "1", "1", "主动 / 资源", "秘密", "1"],
            ["2", "回收利用", "2", "0", "被动 / 资源", "公开", "结束结算 1 次"],
            ["3", "恐吓", "3", "4", "主动 / 控制", "公开", "1"],
            ["4", "绝境", "2", "0", "被动 / 结算", "公开", "手牌开始判定"],
            ["5", "血战", "2", "3", "主动 / 结算", "公开", "1"],
            ["6", "防守", "3", "3", "主动 / 防御", "秘密", "1"],
            ["7", "感知", "3", "0", "被动 / 情报", "秘密", "最多成功 3 次"],
            ["8", "情报", "3", "4", "主动 / 情报", "秘密", "1"],
            ["9", "绝密", "3", "3", "被动 / 防御", "秘密", "自动启动最多 1 次"],
            ["10", "反制", "4", "4", "主动 / 控制", "秘密", "1（仅翻牌前放置）"],
            ["11", "公平", "4", "3", "主动 / 控制", "公开", "1"],
            ["12", "千术", "5", "6", "主动 / 换牌", "混合", "1"],
            ["13", "绝路", "4", "5", "主动 / ALL IN", "公开", "1"],
            ["14", "灵视", "3", "2", "主动 / 元情报", "秘密", "1"],
            ["15", "零化", "5", "6 / 7", "主动 / 牌面", "秘密", "1"],
            ["16", "强运", "5", "改牌 3", "被动 / 幸运", "秘密", "多节点自动"],
            ["17", "天命", "5", "7", "主动 / 牌堆", "秘密", "不限次数（仅转牌）"],
            ["18", "贷款", "2", "2", "主动 / 资源", "混合", "正常 筹 2 / 能 1；受限合计 1"],
            ["19", "警觉", "1", "0", "被动 / 情报", "秘密", "最多成功提示 1 次"],
            ["20", "撤退", "2", "3", "主动 / 防御", "秘密", "1"],
            ["21", "重启", "4", "3", "主动 / 重构", "秘密", "1"],
            ["22", "试探", "1", "2", "主动 / 结算", "秘密", "1"],
            ["23", "伪装", "4", "2", "主动 / 信息控制", "公开", "1"],
            ["24", "终局", "6", "8", "主动 / 超级", "公开", "1"],
        ],
        col_widths=[1.2, 2.2, 1.4, 1.8, 3.2, 1.6, 5.0],
    )
    add_body(doc, "协议均为负载 1、费用 0、被动秘密。摊牌且恰好该牌型并获胜时，标准净收益 ×2。皇家同花顺归同花顺协议。")

    # 3
    add_heading_styled(doc, "3. 主体技能", 1)

    def skill(title, meta, bullets, extra=None):
        add_heading_styled(doc, title, 2)
        add_skill_meta(doc, meta)
        for item in bullets:
            if isinstance(item, tuple):
                kind, text = item
                if kind == "h":
                    add_body(doc, text, bold=True, after=2)
                elif kind == "n":
                    add_bullet(doc, text[1], number=text[0])
                else:
                    add_bullet(doc, text, level=1)
            else:
                add_bullet(doc, item)
        if extra:
            extra()

    skill(
        "1  深呼吸  /  Deep Breath",
        "负载 1  ·  费用 1  ·  主动 / 秘密  ·  每手 1 次",
        [
            "仅自己的合法下注回合，翻牌前至河牌。",
            "支付 1；若直到本手结束没有再发生自己的技能事件，结束时恢复 2。",
            "Fold 后仍结算。无其他技能且自己 Fold：+2，再加败者 +1。",
            "可被反制。若被合法阻止，回收利用按 floor(1 × 50%) = 0 结算，不实际返还能量。",
            "公平清除待恢复状态并抑制本手结束恢复。",
        ],
    )
    skill(
        "2  回收利用  /  Recycle",
        "负载 2  ·  费用 0  ·  被动  ·  每手结束结算 1 次",
        [
            "本手结束时，在「已合法扣费且最终失败」的技能中取原始费用最高的一次，返还 floor(费用 × 50%)。",
            "符合：被反制、被绝密阻止、目标实际不存在或已离开合法区域。",
            "不符合：扣费前非法拒绝、取消、网络重发、反制空放未触发。",
            "不在中途返还。",
        ],
    )
    skill(
        "3  恐吓  /  Intimidation",
        "负载 3  ·  费用 4  ·  主动公开  ·  每手 1 次",
        [
            "前提：双方当前累计标准投入均未超过 500。",
            "本手禁止 Fold；每人累计标准投入最多 500；Bet/Raise 必须让对手仍可合法 Call。",
            "仍允许点 ALL IN，但标准投入最多加到累计 500，并记 allInAction=true；若筹码未全部投入则 stackCommitted=false。",
            "500 上限不管技能倍率造成的最终额外转移。",
            "恐吓下撤退状态可以存在，但不能撤退 Fold。",
            "恐吓下绝路仍可发动，但 Fold ×3 不会发生。",
        ],
    )
    skill(
        "4  绝境  /  Desperation",
        "负载 2  ·  费用 0  ·  被动",
        [
            "每手开始看手牌开始筹码快照：≤ 200 则本手进入绝境。",
            "不按当前筹码、是否 ALL IN、是否落后动态触发。",
            "获胜：标准净收益 ×3，并额外 +1 能量。",
            "平局不触发。公平可清除待结算绝境，并抑制结束 +1。",
        ],
    )
    skill(
        "5  血战  /  Blood Battle",
        "负载 2  ·  费用 3  ·  主动公开  ·  每手 1 次",
        [
            "本手标准净转移 ×2；双方都开则 ×4。",
            "平局不放大。最终不超过败方可支付筹码。",
            "公平可清除。",
        ],
    )
    skill(
        "6  防守  /  Defense",
        "负载 3  ·  费用 3  ·  主动秘密  ·  每手 1 次",
        [
            "自己输掉且不是主动 Fold：最终净损失减半。",
            "赢、平、主动 Fold：无收益，3 费不退。",
            "结算时若确实把公开损失减半，可自然确认「防守」。",
        ],
    )
    skill(
        "7  感知  /  Perception    【FROZEN_V1】",
        "负载 3  ·  费用 0  ·  被动完全秘密",
        [
            "节点：底牌发完、Flop、Turn、River，四次独立判定。",
            "每手最多成功 3 次。",
            "触发概率随自身筹码劣势：均势 25% → 全劣 50%（spec-25-50）。",
            "先选信息类别，再 75% 真 / 25% 假。",
            "假命题必须在当前真实状态下确实为假。",
            "同手避免完全相同、等价或直接逻辑否定。",
            "灵视只能发现「本手发生过感知」，看不到内容与真假。",
            "绝密可阻止感知读取受保护底牌信息。",
            "未解冻前不得改这组概率。",
        ],
    )
    skill(
        "8  情报  /  Intel",
        "负载 3  ·  费用 4  ·  主动秘密  ·  每手 1 次",
        [
            "扣费前锁定模式。",
            "A：系统随机看对手一张底牌，100% 真，不可自选左右；可被绝密阻止，失败后不能改看公共牌。",
            "B：选择任意尚未发出的公共牌位置（含翻牌前看未来 River），100% 真；不受绝密影响。",
        ],
    )
    skill(
        "9  绝密  /  Top Secret",
        "负载 3  ·  费用 3  ·  被动秘密  ·  自动启动最多 1 次",
        [
            "敌方第一次试图读取 / 推断 / 交换 / 零化 / 直接操作本人底牌时，若能量 ≥ 3：自动付 3，阻止该次技能，并持续保护本手剩余时间。",
            "能量不足不发动，不许透支。",
            "阻挡：感知底牌信息、情报底牌、千术对手底牌、零化底牌。",
            "不阻挡：公共牌情报、灵视、纯技能元信息。",
            "第一次真正阻挡时，对方确认「绝密」存在。",
        ],
    )
    skill(
        "10  反制  /  Counter",
        "负载 4  ·  费用 4  ·  主动秘密  ·  仅翻牌前放置  ·  每手 1 次",
        [
            "捕获对手下一次合法主动技能：对方照常扣费 → 技能失败 → 本手不能再主动、也不能再产生新被动事件。",
            "已完成效果不追溯。",
            "可以反制对方正在放置的反制。",
            "非法请求不消耗反制。",
            "公平不能被反制。",
            "整手未触发：结束返还 1（空放净成本 3）。",
            "对天命 100% 有效，不因「牌堆已改」而削弱。",
        ],
    )
    skill(
        "11  公平  /  Fairness",
        "负载 4  ·  费用 3  ·  主动公开  ·  每手 1 次  ·  不能被反制",
        [
            "成功后立即清除双方当前持续 / 预埋 / 待结算 / 尚未完成的技能效果。",
            "包括：反制陷阱、血战、防守、恐吓、绝密当前保护、绝境待结算、深呼吸待恢复、零化、贷款未偿债务、撤退、未触发试探、伪装等。",
            "本手禁止新的主动技能事件与新的被动技能事件。",
            "本手结束恢复全部取消，包括系统败者 +1。",
            "不回滚：已见信息、千术 / 天命 / 重启已完成改牌、已完成筹码直接转移、已完成终局没收。",
            "公平成功后本手不能再发动终局；终局已关闭下注后不再生成普通公平窗口。",
            "仍可清除贷款 pending repayment 与 residual debt，但不退已得筹码 / 能量。",
            "若这次实际清掉任何金额大于 0 的贷款债务：正常信用变为受限；已受限保持受限；已违约则脱离完全封禁并进入受限，不能直接恢复正常。",
            "发动时没有任何贷款债务，不得改变贷款信用状态。",
        ],
    )
    skill(
        "12  千术  /  Cheat",
        "负载 5  ·  费用 6  ·  主动  ·  每手 1 次  ·  可见性混合",
        [
            "用自己一张底牌交换以下目标之一：",
            ("n", (1, "对手指定底牌位置（不知牌值，受绝密）")),
            ("n", (2, "已公开公共牌（改动会自然暴露千术）")),
            ("n", (3, "尚未发出的未来公共牌位")),
            ("n", (4, "下一张有效发牌")),
            ("n", (5, "剩余未发牌堆中均匀随机一张非顶部、且排除下一张有效发牌的牌；旧底牌放回被抽位置")),
            "必须保持 52 张唯一。已完成交换是既定事实。",
        ],
    )
    skill(
        "13  绝路  /  Dead End",
        "负载 4  ·  费用 5  ·  主动公开 ALL IN  ·  每手 1 次",
        [
            "成功后执行当前规则允许的最大合法 ALL IN，并封锁对手本手新的主动 / 被动技能事件。",
            "不清除已有撤退等既定状态。",
            "对手之后普通 Fold：绝路方标准净收益 ×3。",
            "对手 Call 进摊牌：即使绝路方获胜，绝路本身不提供 ×3。",
            "恐吓下仍可发动，投入受 500 截断，但仍记 ALL IN 行为。",
            "绝路 ALL IN 对伪装强制公开。",
            "先绝路后不能新发动撤退；先撤退后绝路仍可撤退 Fold（此时标准净 = 0，Fold ×3 仍为 0）。",
        ],
    )
    skill(
        "14  灵视  /  Clairvoyance",
        "负载 3  ·  费用 2  ·  主动秘密  ·  每手 1 次",
        [
            "私下读取：对手真实能量；本手截至当前已完成的技能事件（含原本秘密技能）。",
            "看不到：感知内容与真假、情报看了哪张、千术隐藏换牌细节、天命点名、强运私人改牌、贷款能量细节、重启抽到的牌、尚未触发的完整构筑。",
            "不能穿透伪装读取 stack / pot / bet / call / contribution。",
            "仍可以知道「伪装这个公开技能已经发生」。",
            "不受绝密阻挡，可被反制。",
        ],
    )
    skill(
        "15  零化  /  Nullification",
        "负载 5  ·  主动完全秘密  ·  每手 1 次  ·  仅 Flop 及之后",
        [
            "公共牌零化费用 6：指定已公开或尚未发出的公共牌位置，最终算牌时对双方视为不存在。",
            "底牌零化费用 7：系统随机对手一张底牌不参与其最佳五张；可被绝密阻止。",
            "双方可零化同一公共牌位置：两次都成功扣费，该牌只失效一次。",
            "持续状态，公平可清除。最终结算时才揭露。",
        ],
    )
    skill(
        "16  强运  /  Fortune    【FROZEN_V1 / soft-v1】",
        "负载 5  ·  被动完全秘密  ·  改牌费用 3",
        [
            "节点：底牌、Flop、Turn、River、手牌结束资源。资源强运不耗费。",
            "底牌已足够优秀则不改底牌；较差时自动改善，玩家不能选择是否触发、换哪张、换成什么。",
            "公共牌强运不得读取对手底牌。",
            "资源强运成功时额外 +1，且不能递归再触发强运。",
            "允许能量降到 -4；再低则该次付费强运不发生。",
            "负能量时除强运外不能有新的主动 / 被动技能事件。",
            "均势、能量 4 时参考触发率：底牌约 8.05%，公共牌约 5.39%，资源 16%。",
            "未解冻前不得改 soft-v1 数值、改牌费用或 -4 下限。",
        ],
    )
    skill(
        "17  天命  /  Destiny",
        "负载 5  ·  费用 7  ·  主动完全秘密  ·  不限次数（仅转牌）",
        [
            "仅 Turn 已公开后的下注阶段、自己的行动回合。",
            "点名一张仍在合法可操作牌堆中的牌，立刻成为下一张真正会发出的 River（有烧牌则对应有效河牌位）。",
            "立即完成的牌堆修改，公平不能回滚。",
            "目标非法：7 费照付且失败。",
            "反制 100% 阻止。对手仍可 Fold，7 费不退。",
            "携带者能量上限 8→10，初始仍为 4；公平不能改回上限。",
            "本技能不设每手次数上限，但受阶段与能量约束。",
        ],
    )

    add_heading_styled(doc, "18  贷款  /  Loan", 2)
    add_skill_meta(doc, "负载 2  ·  费用 2  ·  主动  ·  可见性混合")
    for t in [
        "扣费前锁定分支。",
        "两种贷款可以同时存在。",
        "信用状态三档，额度来自当前信用，不跨 Match。",
    ]:
        add_bullet(doc, t)
    add_table(
        doc,
        ["信用", "筹码贷 / 手", "能量贷 / 手", "合计 / 手"],
        [
            ["NORMAL_CREDIT 正常", "2", "1", "3"],
            ["RESTRICTED_CREDIT 受限", "1", "1", "1（整手只能选一个模式发动一次）"],
            ["DEFAULTED 违约", "0", "0", "0（整技能不可发动）"],
        ],
        col_widths=[4.8, 2.8, 2.8, 6.0],
    )
    for t in [
        "任意偿还失败形成的 chipDebt 或 energyDebt 未清偿时，进入 DEFAULTED，整技能封禁。",
        "只有从玩家筹码 / 能量中真实扣除才算偿还。公平清除、比赛结束失效、调试重置、状态覆盖、其他技能免除都不算实际偿还。",
    ]:
        add_bullet(doc, t)
    add_body(doc, "筹码贷款（公开）", bold=True, after=2)
    for t in [
        "每次从对手取得 100（不足则取光并斩杀）。",
        "下一手结束偿还 150；两笔按各自到期分别偿还。",
        "偿还使自己归 0 也可斩杀自己。",
        "已拿到的筹码不因公平或赛后失效而退回。",
    ]:
        add_bullet(doc, t)
    add_body(doc, "能量贷款（秘密）", bold=True, after=2)
    for t in [
        "支付 2 后立即 +5，不超过个人上限。",
        "下一手结束还 6；不足先扣当前能量，剩余形成能量债务。",
        "之后任何能量收入优先偿债。",
    ]:
        add_bullet(doc, t)
    add_body(doc, "信用转移", bold=True, after=2)
    for t in [
        "任何已存在的 pending / residual 贷款债务金额 > 0 被公平或其他非真实偿还机制清除 → RESTRICTED_CREDIT。",
        "已受限时再次被公平清债 → 保持受限，不能洗回正常。",
        "到期无法全额真实偿还并形成 residual > 0 → DEFAULTED。",
        "受限后恢复正常，必须同时满足：这笔贷款在受限状态下新产生；正常到期；到期时全部用真实资源支付；没有任何部分被免除；最终 residual == 0。部分真实偿还后剩余被公平清除，不算恢复。",
        "违约后 residual 被真实全部还清 → NORMAL_CREDIT。",
        "违约后 residual 被公平清除 → RESTRICTED_CREDIT，贷款重新可用，但不是洗白。",
        "反制抓住贷款：费用照付，不获得资源，不生成偿还或 residual，不改变信用。",
        "整场比赛合法结束后：未到期偿还、residual、信用状态全部清空；下一场从正常信用开始。",
    ]:
        add_bullet(doc, t)
    add_body(doc, "赛后", bold=True, after=2)
    for t in [
        "必须先摊牌、分配底池、更新筹码，再判断比赛是否结束。",
        "比赛已依法结束：整房所有未到期 / 未清偿贷款与能量债务直接失效。",
        "不得赛后再还款、逆转胜负、复活败者。",
        "伪装下公开贷款只说「发动贷款」，不公布拿了多少、欠多少或斩杀数字。",
    ]:
        add_bullet(doc, t)

    skill(
        "19  警觉  /  Alert",
        "负载 1  ·  费用 0  ·  被动完全秘密",
        [
            "只监听敌方合法、正式提交且仍对自己隐藏的主动技能事件。",
            "公开主动、被动、非法、取消、重发都不触发。",
            "概率：10% → 25% → 40% → 55% → 70% → 85% → 100%。",
            "漏过升一级，成功后重置为 10%；敏锐度跨手保留。",
            "成功后不在技能瞬间提示，而在自己下一次合法下注决策开始时显示：「你隐约察觉到对手似乎进行了秘密行动。」",
            "不透露名称 / 类型 / 数量 / 费用 / 结果 / 精确时间。本手最多成功提示一次。",
            "伪装是公开主动，绝不触发警觉。",
        ],
    )

    add_heading_styled(doc, "20  撤退  /  Retreat", 2)
    add_skill_meta(doc, "负载 2  ·  费用 3  ·  主动秘密防御  ·  每手 1 次")
    for t in [
        "在自己任意合法下注窗口支付 3，秘密建立撤退状态。",
        "允许同一窗口立即 Fold，也允许稍后在其他合法 Fold 窗口使用。",
        "不要求提前一街、最小延迟、冷却。",
    ]:
        add_bullet(doc, t)
    add_body(doc, "撤退 Fold", bold=True, after=2)
    for i, t in enumerate(
        [
            "手牌结束。",
            "双方本手全部标准贡献原路返还（含盲注与标准 ALL IN 贡献）。",
            "标准净转移 = 0。",
            "使用者不获得败者 +1，因此净技能成本就是 3。",
        ],
        1,
    ):
        add_bullet(doc, t, number=i)
    add_body(doc, "不回滚：贷款 / 终局直接筹码、已付能量、已见信息、已完成换牌。")
    add_body(doc, "未使用撤退 Fold", bold=True, after=2)
    for t in ["正常赢：不回能量，净 -3。", "正常输：仍有败者 +1，净 -2。"]:
        add_bullet(doc, t)
    add_body(doc, "被反制：3 已付，状态不建立，不自动 Fold，留在原窗口；有回收则可按规则返 1。")
    add_body(doc, "交互", bold=True, after=2)
    for t in [
        "先撤退后绝路：仍可撤退 Fold；0 × 3 = 0。",
        "先绝路：不能新发动撤退。",
        "公平清除状态，3 费不退。",
        "恐吓下不能 Fold。",
        "终局关闭下注后状态可仍在，但没有 Fold 窗口；终局不主动删除撤退。",
        "撤退 Fold 不触发试探。",
    ]:
        add_bullet(doc, t)

    skill(
        "21  重启  /  Restart",
        "负载 4  ·  费用 3  ·  主动秘密  ·  每手 1 次",
        [
            "两张底牌一起洗回合法剩余牌堆，再随机抽两张。",
            "允许抽回原来那两张。完成即既定事实。",
            "必须保持 52 张唯一。",
        ],
    )
    skill(
        "22  试探  /  Probe",
        "负载 1  ·  费用 2  ·  主动秘密  ·  每手 1 次",
        [
            "对手之后主动选择普通 Fold：基础标准净收益 +50，再算倍率。",
            "例：基础 100、试探成功、血战 → (100+50)×2 = 300。",
            "自己 Fold 不触发自己的试探。",
            "撤退 Fold 不是普通 Fold。",
            "未触发状态可被公平清除。",
        ],
    )

    add_heading_styled(doc, "23  伪装  /  Disguise", 2)
    add_skill_meta(doc, "负载 4  ·  费用 2  ·  主动公开  ·  每手 1 次")
    add_body(doc, "隐藏的是「受影响玩家能获得的筹码数值信息」，必须从玩家专属视图模型 / 序列化层裁剪，不能只靠 CSS。")
    add_body(doc, "单方：发动者视角正常。受影响方看不到：", bold=True, after=2)
    for t in [
        "双方剩余筹码、筹码变化、本手累计投入。",
        "Pot、Bet / Raise / Call 金额、最小加注、需补多少。",
        "「余额不足 XXX」等可反推筹码的数字。",
        "Call 按钮只显示「Call」；Raise 不显示 Minimum Raise 数字。",
    ]:
        add_bullet(doc, t)
    add_body(doc, "超额输入：受影响方按真实最大可投入自动封顶为合法 ALL IN，不返回筹码不足数字。本人立刻知道自己是否 ALL IN，但不显示「你剩余 / 实际投入 XXX」。")
    add_body(doc, "普通 ALL IN：对受影响的对手伪装成普通 Bet / Raise / Call，无 ALL IN 标签 / 动画 / 特效。本人 ALL IN 状态始终可见，包括双重伪装。绝路 ALL IN 强制对双方公开。")
    add_body(doc, "双方伪装：进入黑暗筹码状态，仍各自知道自己的 ALL IN。")
    for t in [
        "灵视不能看具体筹码数字。公开技能不触发警觉。",
        "公平后恢复未来显示，不回填已经隐藏时刻的历史数值日志。",
        "摊牌 / 结算画面在伪装持续期间仍隐藏 Pot 与筹码变化，只公开谁赢、牌型、手牌。",
    ]:
        add_bullet(doc, t)

    add_heading_styled(doc, "24  终局  /  Endgame", 2)
    add_skill_meta(doc, "负载 6  ·  费用 8  ·  主动公开超级技能  ·  每手 1 次")
    for t in [
        "可在自己的合法下注窗口使用，包括面对 Bet、Raise、普通 ALL IN。",
        "另有 Call-to-zero 专属窗口：对手合法 Call 后真实剩余筹码 == 0，且下一张公共牌尚未发出、Showdown 尚未执行。",
        "该窗口只给刚刚造成 Call-to-zero 的进攻方。",
        "界面必须提供「发动终局 / 放弃」；超时视为放弃。",
        "Bot 按自身可见牌力判断，不得无脑自动发动。",
    ]:
        add_bullet(doc, t)
    add_body(doc, "处决资格：仅当对手因真实标准 ALL IN 或 Call-to-zero 导致 remaining chips == 0。恐吓伪 ALL IN 若真实仍有筹码，则 execution = false。不能把终局没收之后才变成 0 再伪造处决。")
    add_body(doc, "严格顺序", bold=True, after=2)
    for i, t in enumerate(
        [
            "确认合法",
            "支付 8",
            "处理反制。命中则失败：不没收、不锁池、不关下注、不处决、不自动发牌；有回收可返 4",
            "成功公开",
            "计算双方标准贡献，matched = min",
            "仅没收对手未匹配部分，类型为直接技能筹码转移，不进倍率",
            "匹配部分锁定，双方不能再 Bet / Raise / Call / Fold",
            "记录 execution（依据发动前真实剩余）",
            "按当前真实牌堆发完尚未发出的公共牌",
            "Showdown",
            "比较最佳五张",
        ],
        1,
    ):
        add_bullet(doc, t, number=i)
    add_body(doc, "处决摊牌", bold=True, after=2)
    for t in [
        "execution = false：完整普通比牌。",
        "execution = true 且牌型等级不同：正常比较。皇家同花顺 10 > 同花顺 9。",
        "execution = true 且等级完全相同：发动者直接胜，不再比内部大小 / kicker。",
        "专属斩杀特效只在「处决覆盖了普通内部比较」时标记 endgameExecutionOverride。",
        "本来就靠更高牌型赢，不得标斩杀。",
        "发动者自己多出的未匹配注按普通德州退还，不进没收池。",
        "例：matched 基础 100，终局没收 700，血战 ×2 → 标准 200 + 直接 700 = 900，禁止 (100+700)×2。",
    ]:
        add_bullet(doc, t)

    # 4
    add_heading_styled(doc, "4. 协议（9）", 1)
    add_body(doc, "均为：负载 1 / 费用 0 / 被动秘密 / 结算。")
    add_body(doc, "触发条件", bold=True, after=2)
    for i, t in enumerate(
        [
            "必须进入 Showdown",
            "最终最佳五张恰好为指定牌型",
            "自己赢得本手",
            "Fold 获胜不触发",
            "若本手已有自己技能产生的其他筹码倍率，本协议不触发",
            "对手技能倍率不阻止本协议，可继续乘法叠加",
            "奖励：标准净收益 ×2",
        ],
        1,
    ):
        add_bullet(doc, t, number=i)
    add_table(
        doc,
        ["协议", "对应牌型"],
        [
            ["协议--高牌", "High Card"],
            ["协议--对子", "One Pair"],
            ["协议--两对", "Two Pair"],
            ["协议--三条", "Trips"],
            ["协议--顺子", "Straight"],
            ["协议--同花", "Flush"],
            ["协议--葫芦", "Full House"],
            ["协议--四条", "Quads"],
            ["协议--同花顺", "Straight Flush；皇家同花顺也走本协议"],
        ],
        col_widths=[4.8, 11.6],
    )
    add_body(doc, "终局处决比牌仍把皇家同花顺与同花顺视为不同等级（10 与 9）。协议层把两者都算同花顺。")

    # 5
    add_heading_styled(doc, "5. 关键交互摘要", 1)
    add_table(
        doc,
        ["交互", "现行结论"],
        [
            ["公平 × 反制", "公平免疫反制"],
            ["公平 × 贷款", "清未偿债务，不退已得资源；清掉金额大于 0 的债会使贷款进入或保持信用受限，不能把违约直接洗回正常信用"],
            ["公平 × 撤退", "清状态，3 费不退；之后不能新撤退"],
            ["公平 × 伪装", "清状态，恢复未来视图，不回填历史数字"],
            ["公平 × 终局", "公平已成功则不能发动终局；终局已关下注则不再给普通公平窗口；已没收不回滚"],
            ["公平 × 零化", "零化是持续状态，可被清除"],
            ["伪装 × 灵视", "知道伪装发生，看不到筹码数字"],
            ["伪装 × 警觉", "公开主动，不触发"],
            ["伪装 × 贷款", "公开贷款不带数字"],
            ["伪装 × 绝路", "绝路 ALL IN 强制公开"],
            ["撤退 × 绝路", "先撤后绝仍可撤；先绝后不能新撤"],
            ["撤退 × 试探", "撤退 Fold 不触发试探"],
            ["撤退 × 恐吓", "不能 Fold"],
            ["撤退 × 终局", "无 Fold 窗口，状态不必删除"],
            ["终局 × 反制", "先付 8 再处理反制；命中则无没收、无处决"],
            ["终局 × 血战等倍率", "直接没收不进倍率"],
            ["终局 × 协议", "处决同级由发动者胜；协议仍看最终牌型是否恰好匹配"],
        ],
        col_widths=[4.2, 12.2],
    )

    # 6
    add_heading_styled(doc, "6. 明确不回滚的事实", 1)
    add_body(doc, "以下在公平、撤退、反制失败补偿中都不是「把时间倒回去」：")
    for t in [
        "已经看到的信息",
        "千术 / 重启已完成换牌",
        "天命已完成牌堆修改",
        "贷款已取得的筹码或能量",
        "终局已完成的直接没收",
        "已支付且未被回收规则覆盖的能量",
    ]:
        add_bullet(doc, t)

    # 7
    add_heading_styled(doc, "7. 冻结声明", 1)
    add_table(
        doc,
        ["系统", "状态", "变体", "冻结日"],
        [
            ["感知", "FROZEN_V1", "spec-25-50", "2026-08-20"],
            ["强运", "FROZEN_V1", "soft-v1", "2026-08-20"],
            ["贷款信用", "正式规则", "信用受限 V2", "2026-08-21"],
            ["贷款数值", "现行冻结", "100 / 150、费用 2、正常额度 2+1", "2026-08-21"],
            ["公平数值", "现行冻结", "负载 4 / 费用 3", "2026-08-21"],
        ],
        col_widths=[3.2, 4.4, 4.4, 4.4],
    )
    add_callout(doc, "解冻前禁止改：感知触发 / 真假概率、强运 soft-v1 数值、改牌费用 3、能量下限 -4、贷款 100 / 150 与费用 2、公平负载 4 / 费用 3。贷款信用受限已升为正式规则，暂不继续削弱。")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(OUT_PATH)


if __name__ == "__main__":
    build()
